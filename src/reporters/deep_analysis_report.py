"""
深度分析报告生成器
生成 AI 资讯的深度分析 Word 报告
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from collections import Counter
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import io
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DeepAnalysisReport:
    """深度分析报告生成器"""
    
    def __init__(self, news_list):
        self.news_list = news_list
        self.report_date = datetime.now()
        
    def generate(self, output_file="AI_深度分析报告.docx"):
        """生成完整的深度分析报告"""
        logger.info(f"正在生成深度分析报告...")
        
        doc = Document()
        
        # 1. 标题
        self._add_title(doc)
        
        # 2. 数据概览
        self._add_data_overview(doc)
        
        # 3. 热点话题分析
        self._add_hot_topics(doc)
        
        # 4. 重要资讯解读
        self._add_key_news_analysis(doc)
        
        # 5. 趋势分析
        self._add_trend_analysis(doc)
        
        # 6. 行动建议
        self._add_actionable_recommendations(doc)
        
        # 7. 附件：今日资讯列表
        self._append_news_list(doc)
        
        # 保存文档
        doc.save(output_file)
        logger.info(f"✓ 报告已生成: {output_file}")
        
        return output_file
    
    def _add_title(self, doc):
        """添加报告标题"""
        title = doc.add_heading()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = title.add_run(f"AI 行业深度分析报告")
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.bold = True
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"报告日期：{self.report_date.strftime('%Y年%m月%d日')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_data_overview(self, doc):
        """添加数据概览"""
        doc.add_heading('一、数据概览', level=1)
        
        # 统计数据
        total_count = len(self.news_list)
        source_stats = Counter([n['source'] for n in self.news_list])
        category_stats = Counter([n.get('sub_category', n.get('category', 'Other')) for n in self.news_list])
        
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 标题行
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        # 添加数据
        data = [
            ('今日采集总数', f"{total_count} 条"),
            ('资讯来源数', f"{len(source_stats)} 个"),
            ('主要来源', ', '.join([f"{k}({v})" for k, v in source_stats.most_common(3)])),
            ('热门分类', ', '.join([f"{k}({v})" for k, v in category_stats.most_common(3)]))
        ]
        
        for item, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = value
        
        doc.add_paragraph()  # 空行
    
    def _add_hot_topics(self, doc):
        """添加热点话题分析"""
        doc.add_heading('二、热点话题分析', level=1)
        
        # 提取关键词
        all_keywords = []
        for news in self.news_list[:15]:
            title = news['title'].lower()
            words = [word.strip(' ,.!?;:"\'') for word in title.split()]
            all_keywords.extend([w for w in words if len(w) > 2])
        
        keyword_freq = Counter(all_keywords)
        top_keywords = keyword_freq.most_common(15)
        
        # 添加文本分析
        p = doc.add_paragraph()
        run = p.add_run('今日热点关键词：')
        run.bold = True
        run.font.size = Pt(12)
        
        for keyword, count in top_keywords:
            p.add_run(f'{keyword}({count}) ')
            p.add_run('· ')
        
        doc.add_paragraph()
        
        # 分析热点话题
        doc.add_heading('热点话题解读', level=2)
        
        topics = self._analyze_topics(top_keywords)
        for topic in topics:
            doc.add_paragraph(topic, style='List Bullet')
        
        doc.add_paragraph()
    
    def _analyze_topics(self, keywords):
        """分析热点话题"""
        topics = []
        
        # 基于 Top 10 关键词生成话题分析
        tech_keywords = ['llm', 'gpt', '大模型', 'agent', '智能体', 'openai', 'deepseek']
        company_keywords = ['google', 'microsoft', 'meta', 'apple', 'nvidia']
        application_keywords = ['应用', '落地', '场景', '产品', '发布']
        
        tech_found = [k for k, v in keywords if k in tech_keywords]
        company_found = [k for k, v in keywords if k in company_keywords]
        app_found = [k for k, v in keywords if k in application_keywords]
        
        if tech_found:
            topics.append(f"**技术焦点**：{', '.join(tech_found[:3])} 相关资讯占据今日主流，显示行业内技术突破和模型更新的高度关注。")
        
        if company_found:
            topics.append(f"**企业动态**：{', '.join(company_found[:3])} 等科技巨头的动向引发广泛关注，反映市场竞争格局的变化。")
        
        if app_found:
            topics.append(f"**应用趋势**：{', '.join(app_found[:3])} 等应用场景的讨论增多，表明 AI 技术正在加速落地和商业化。")
        
        if not topics:
            topics.append("今日资讯主题分散，关注点多元化，建议关注后续趋势。")
        
        return topics
    
    def _add_key_news_analysis(self, doc):
        """添加重要资讯解读"""
        doc.add_heading('三、重要资讯深度解读', level=1)
        
        # 取 Top 5 资讯进行深度解读
        top_news = self.news_list[:5]
        
        for i, news in enumerate(top_news, 1):
            doc.add_heading(f'{i}. {news["title"][:40]}...', level=2)
            
            # 来源和评分
            p = doc.add_paragraph()
            p.add_run(f'来源：{news["source"]}\n')
            p.add_run(f'评分：{news.get("score", 0)}\n')
            p.add_run(f'分类：{news.get("sub_category", news.get("category", "未分类"))}')
            
            # 深度解读
            doc.add_heading('深度解读', level=3)
            analysis = self._analyze_single_news(news)
            doc.add_paragraph(analysis)
            
            # 行业影响
            doc.add_heading('行业影响', level=3)
            impact = self._analyze_impact(news)
            doc.add_paragraph(impact)
            
            doc.add_paragraph()  # 空行
    
    def _analyze_single_news(self, news):
        """分析单条资讯"""
        analysis = f"该资讯涉及{news.get('sub_category', 'AI')}领域，"
        
        # 基于标题和摘要生成分析
        title = news['title'].lower()
        summary = news.get('summary', '').lower()
        
        if '发布' in title or '推出' in title:
            analysis += "展现了新的产品或技术发布，体现了行业在相关方向的持续投入和创新动力。"
        elif '融资' in title or '投资' in title:
            analysis += "反映了资本市场对该领域的关注和信心，可能预示着相关技术和应用的加速发展。"
        elif '突破' in title or '创新' in title:
            analysis += "代表了技术层面的重要进展，可能对行业格局产生深远影响。"
        elif '合作' in title or '联盟' in title:
            analysis += "表明产业链上下游的合作趋势，有助于推动技术标准化和规模化应用。"
        else:
            analysis += "具有重要的信息价值，建议持续关注相关进展和后续报道。"
        
        return analysis
    
    def _analyze_impact(self, news):
        """分析行业影响"""
        impact = "短期影响："
        
        score = news.get('score', 0)
        if score >= 15:
            impact += "该资讯可能引发行业广泛关注和讨论。"
        elif score >= 10:
            impact += "该资讯在特定领域具有重要参考价值。"
        else:
            impact += "该资讯提供了一定的行业信息参考。"
        
        impact += "\n\n长期影响："
        impact += "建议持续关注该领域的发展趋势，评估对个人/业务策略的潜在影响。"
        
        return impact
    
    def _add_trend_analysis(self, doc):
        """添加趋势分析"""
        doc.add_heading('四、趋势分析', level=1)
        
        # 基于关键词和分类分析趋势
        category_stats = Counter([n.get('sub_category', n.get('category', 'Other')) for n in self.news_list])
        
        # 趋势1：技术演进
        doc.add_heading('技术演进趋势', level=2)
        
        if 'LLM' in category_stats:
            doc.add_paragraph('• 大语言模型（LLM）依然是行业焦点，多模态、效率优化和成本降低是主要发展方向。')
        
        if 'Agent' in category_stats:
            doc.add_paragraph('• AI Agent 智能体概念持续升温，自动化、自主决策能力成为研究热点。')
        
        if 'AI Tool' in category_stats:
            doc.add_paragraph('• AI 工具和应用场景不断丰富，从实验室走向实际应用的趋势明显。')
        
        # 趋势2：应用落地
        doc.add_heading('应用落地趋势', level=2)
        doc.add_paragraph('• AI 技术正在加速向各行各业渗透，金融、医疗、教育等领域的应用案例增多。')
        doc.add_paragraph('• 企业级 AI 解决方案需求增长，定制化、垂直化成为竞争焦点。')
        
        # 趋势3：产业生态
        doc.add_heading('产业生态趋势', level=2)
        doc.add_paragraph('• 科技巨头、创业公司、传统企业都在布局 AI，产业生态日趋多元化。')
        doc.add_paragraph('• 开源社区和开发者生态的重要性凸显，推动技术普及和创新。')
        
        doc.add_paragraph()
    
    def _add_actionable_recommendations(self, doc):
        """添加行动建议"""
        doc.add_heading('五、行动建议', level=1)
        
        recommendations = [
            {
                'title': '学习建议',
                'items': [
                    '关注今日 Top 3 资讯的详细报道，深入了解相关技术原理',
                    '学习推荐的热点技术（如 LLM、Agent）的核心概念和应用场景',
                    '跟踪相关公司的产品发布和动态，了解行业最新趋势'
                ]
            },
            {
                'title': '实践建议',
                'items': [
                    '尝试使用今日资讯中提到的 AI 工具或产品，体验其实际效果',
                    '基于热点技术方向，思考如何应用到自己的工作或项目中',
                    '参与相关技术社区或论坛，与同行交流观点和经验'
                ]
            },
            {
                'title': '规划建议',
                'items': [
                    '评估今日资讯对未来 3-6 个月行业发展的潜在影响',
                    '调整个人或团队的 AI 技术学习计划，重点关注热点领域',
                    '考虑将相关技术纳入产品规划或业务拓展方向'
                ]
            }
        ]
        
        for rec in recommendations:
            doc.add_heading(rec['title'], level=2)
            for item in rec['items']:
                doc.add_paragraph(item, style='List Bullet')
            doc.add_paragraph()
    
    def _append_news_list(self, doc):
        """添加完整的资讯列表"""
        doc.add_heading('附件：今日资讯完整列表', level=1)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 标题行
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '标题'
        hdr_cells[1].text = '来源'
        hdr_cells[2].text = '链接'
        
        # 设置列宽
        table.columns[0].width = Inches(4.0)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(2.0)
        
        # 添加资讯
        for news in self.news_list:
            row_cells = table.add_row().cells
            row_cells[0].text = news['title'][:50] + '...' if len(news['title']) > 50 else news['title']
            row_cells[1].text = news['source']
            row_cells[2].text = news['link']
        
        # 添加免责声明
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        run = disclaimer.add_run('免责声明：本报告基于公开资讯生成，仅供参考，不构成投资建议。')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True


if __name__ == "__main__":
    # 测试代码
    test_news = [
        {
            'title': 'OpenAI 发布 GPT-5 预览版',
            'link': 'https://example.com/gpt5',
            'summary': 'OpenAI 今天宣布发布 GPT-5 预览版，带来了新的多模态能力和更强的推理能力。',
            'source': '量子位',
            'category': 'Chinese',
            'sub_category': 'LLM',
            'score': 15,
            'collected_at': datetime.now()
        },
        {
            'title': 'DeepSeek-R1 开源代码库突破 10k Stars',
            'link': 'https://example.com/deepseek',
            'summary': 'DeepSeek-R1 开源项目在 GitHub 上获得了超过 10k Stars，显示了社区对该项目的高度关注。',
            'source': 'MIT Technology Review',
            'category': 'Media',
            'sub_category': 'AI Tool',
            'score': 12,
            'collected_at': datetime.now()
        }
    ]
    
    generator = DeepAnalysisReport(test_news)
    report_file = generator.generate("测试报告.docx")
    print(f"测试报告已生成: {report_file}")
